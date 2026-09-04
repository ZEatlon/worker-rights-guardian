"""Word 文档举报材料生成器（可选依赖 python-docx）。

仅在安装了 ``python-docx`` 的情况下可用，缺失时 ``WordReportGenerator``
构造会抛 ``ImportError``，调用方应自行捕获并提示安装::

    pip install python-docx
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any


class WordReportGenerator:
    """生成 Word 格式的举报材料。"""

    def __init__(self, output_dir: str | Path | None = None) -> None:
        try:
            from docx import Document  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            from docx.shared import Pt  # type: ignore
        except ImportError as exc:
            raise ImportError(
                "生成 Word 报告需要 python-docx，请先运行 "
                "`pip install python-docx`。"
            ) from exc

        self._Document = Document
        self._WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
        self._Pt = Pt

        if output_dir is None:
            output_dir = Path.cwd() / "reports"
        self.output_dir = Path(output_dir).expanduser()
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        case_data: dict[str, Any],
        evidence_list: list[dict[str, Any]],
        case_name: str,
        *,
        lang: str = "zh",
    ) -> str:
        """生成 Word 报告并返回路径。

        Args:
            case_data: 案件信息（含 violations）。
            evidence_list: 证据项字典列表（可来自 ``EvidenceItem.to_dict()``）。
            case_name: 用于文件名。
            lang: 'zh' / 'en' / 'ja' / 'ko'，仅影响标题语言。
        """
        doc = self._Document()

        # 默认字体
        try:
            style = doc.styles["Normal"]
            style.font.name = "宋体"
            style.font.size = self._Pt(12)
        except (KeyError, AttributeError):
            pass

        title_text = {
            "zh": "劳动监察投诉举报材料",
            "en": "Labor Rights Complaint Report",
            "ja": "労働基準監督告発文書",
            "ko": "근로기준 감독 신고 서류",
        }.get(lang, "Labor Rights Complaint Report")

        # 标题
        title = doc.add_heading(title_text, level=0)
        for run in title.runs:
            run.font.size = self._Pt(20)
        title.alignment = self._WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息表
        doc.add_heading(self._t("basic_info", lang), level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        rows = self._basic_rows(case_data, lang)
        for label, value in rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        # 投诉事项
        doc.add_heading(self._t("violations", lang), level=1)
        violations = case_data.get("violations") or []
        if violations:
            for i, v in enumerate(violations, 1):
                doc.add_heading(
                    f"[{i}] {v.get('type', '')}", level=2
                )
                doc.add_paragraph(
                    self._t("description", lang)
                    + ": " + str(v.get("description", ""))
                )
                if v.get("amount"):
                    doc.add_paragraph(
                        self._t("amount", lang)
                        + ": " + str(v.get("amount", ""))
                    )
        else:
            doc.add_paragraph(self._t("no_violations", lang))

        # 证据清单
        doc.add_heading(self._t("evidence_list", lang), level=1)
        if evidence_list:
            et = doc.add_table(rows=1, cols=3)
            et.style = "Light Grid Accent 1"
            et.add_row().cells[0].text = "#"
            et.add_row().cells[1].text = self._t("evidence_name", lang)
            et.add_row().cells[2].text = self._t("evidence_type", lang)
            for i, ev in enumerate(evidence_list, 1):
                row = et.add_row().cells
                row[0].text = str(i)
                row[1].text = str(ev.get("name", ""))
                row[2].text = str(ev.get("file_type", ""))
        else:
            doc.add_paragraph(self._t("no_evidence", lang))

        # 真实性声明
        doc.add_paragraph()
        declaration = {
            "zh": "以上材料属实，本人愿承担相应法律责任。",
            "en": (
                "I declare that the information provided above is true to the "
                "best of my knowledge, and I am willing to bear the "
                "corresponding legal responsibility."
            ),
            "ja": (
                "上記内容に虚偽なく、私が法的責任を負うことを誓約いたします。"
            ),
            "ko": (
                "위 내용이 사실과 다름없으며, 본인은 관련 법적 책임을 부담할 "
                "것을 서약합니다."
            ),
        }.get(lang, "")
        if declaration:
            p = doc.add_paragraph(declaration)
            p.alignment = self._WD_ALIGN_PARAGRAPH.CENTER

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = self._safe(case_name) or "case"
        filepath = self.output_dir / f"report_{safe_name}_{ts}.docx"
        doc.save(filepath)
        return str(filepath)

    # ---------- 辅助 ----------

    @staticmethod
    def _safe(name: str) -> str:
        if not name:
            return ""
        keep = "-_.()"
        return "".join(
            ch if (ch.isalnum() or ch in keep) else "_" for ch in name
        )[:80]

    @staticmethod
    def _t(key: str, lang: str) -> str:
        messages: dict[str, dict[str, str]] = {
            "basic_info": {
                "zh": "一、基本信息",
                "en": "1. Basic Information",
                "ja": "一、基本情報",
                "ko": "1. 기본 정보",
            },
            "violations": {
                "zh": "二、投诉事项",
                "en": "2. Violations",
                "ja": "二、違反事実",
                "ko": "2. 위반 사항",
            },
            "description": {
                "zh": "详细情况",
                "en": "Description",
                "ja": "詳細",
                "ko": "상세",
            },
            "amount": {
                "zh": "涉及金额/时间",
                "en": "Amount/Period",
                "ja": "金額/期間",
                "ko": "금액/period",
            },
            "no_violations": {
                "zh": "（无）",
                "en": "(none)",
                "ja": "（なし）",
                "ko": "（없음）",
            },
            "evidence_list": {
                "zh": "三、证据清单",
                "en": "3. Evidence List",
                "ja": "三、証拠一覧",
                "ko": "3. 증거 목록",
            },
            "evidence_name": {
                "zh": "名称",
                "en": "Name",
                "ja": "名称",
                "ko": "이름",
            },
            "evidence_type": {
                "zh": "类型",
                "en": "Type",
                "ja": "タイプ",
                "ko": "유형",
            },
            "no_evidence": {
                "zh": "（无证据）",
                "en": "(no evidence)",
                "ja": "（証拠なし）",
                "ko": "（증거 없음）",
            },
        }
        return messages.get(key, {}).get(lang, messages.get(key, {}).get("en", key))

    def _basic_rows(self, case_data: dict[str, Any], lang: str) -> list[tuple[str, str]]:
        today = datetime.now().strftime("%Y-%m-%d")
        labels = {
            "complainant": {"zh": "投诉人", "en": "Complainant", "ja": "告発者", "ko": "신고자"},
            "respondent": {"zh": "被投诉单位", "en": "Respondent", "ja": "対象企業", "ko": "피신고 기업"},
            "address": {"zh": "单位地址", "en": "Address", "ja": "住所", "ko": "주소"},
            "phone": {"zh": "联系电话", "en": "Phone", "ja": "電話", "ko": "전화"},
            "email": {"zh": "电子邮箱", "en": "Email", "ja": "メール", "ko": "이메일"},
            "date": {"zh": "出具日期", "en": "Report date", "ja": "作成日", "ko": "작성일"},
        }

        rows = [
            (labels["complainant"][lang],
             case_data.get("worker_name", "")),
            (labels["respondent"][lang],
             case_data.get("company_name", "")),
            (labels["address"][lang],
             case_data.get("company_address", "")),
            (labels["phone"][lang],
             case_data.get("worker_phone", "")),
            (labels["email"][lang],
             case_data.get("worker_email", "")),
            (labels["date"][lang], today),
        ]
        return rows


__all__ = ["WordReportGenerator"]